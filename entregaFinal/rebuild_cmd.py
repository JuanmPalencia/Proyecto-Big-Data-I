# -*- coding: utf-8 -*-
"""
GTP_Entrega_Modulo_CMD_v2.docx — version final limpia
Estructura:
  1. Resumen Ejecutivo
  2. Idea de Negocio y Plataforma Web
  3. KPIs y OKRs
  4. Estructura Analitica
  5. Cuadros de Mando
      5.1 CMDs Operativos internos (CMD1 Mapa Inversion + CMD2 EKC) — CMD1 con mockup, CMD2 teorico
      5.2 CMDs de Cliente (CMD3 Ambiental + CMD4 Financiero) — CMD4 con mockup, CMD3 teorico
          Nota: mapa europeo pendiente de ubicacion (vista separada o componente de CMD4)
Sin referencias a API, sin secciones de datos/ETL/modelos/arquitectura tecnica.
"""

import zipfile, os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "C:/Users/Juanm/Big Data I/Proyecto-Big-Data-I/entregaFinal/CMD - Planficacion tarea/"
SRC  = BASE + "GTP_Entrega_Modulo_CMD_v2.docx"
DST  = BASE + "GTP_Entrega_Modulo_CMD_v2.docx"
IMG  = BASE + "_tmp_image1.png"

C_GREEN_DARK  = "2E7D32"
C_GREEN_MID   = "388E3C"
C_GREEN_LIGHT = "C8E6C9"
C_GREEN_TEXT  = "1B5E20"
C_GRAY        = "757575"
C_AMBER       = "F57F17"   # advertencias/notas

with zipfile.ZipFile(SRC, "r") as z:
    with z.open("word/media/image1.png") as s, open(IMG, "wb") as d:
        d.write(s.read())

doc = Document()
for sec in doc.sections:
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(2.5)

# ── helpers de estilo ─────────────────────────────────────────────────────────

def rgb(h):
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def valign_center(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    v = OxmlElement("w:vAlign"); v.set(qn("w:val"), "center"); tcPr.append(v)

def hrule(color="2E7D32", sz=16):
    par = doc.add_paragraph()
    pPr = par._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), str(sz))
    bot.set(qn("w:space"), "1");    bot.set(qn("w:color"), color)
    pBdr.append(bot); pPr.append(pBdr)

def add_page_numbers():
    for section in doc.sections:
        footer = section.footer
        par    = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER; par.clear()
        run = par.add_run()
        for tag, attrs in [
            ("w:fldChar",   {"w:fldCharType": "begin"}),
            ("w:instrText", None),
            ("w:fldChar",   {"w:fldCharType": "separate"}),
            ("w:fldChar",   {"w:fldCharType": "end"}),
        ]:
            el = OxmlElement(tag)
            if attrs is None:
                el.set(qn("xml:space"), "preserve"); el.text = "PAGE"
            else:
                for k, v in attrs.items(): el.set(qn(k), v)
            run._r.append(el)

def toc():
    par = doc.add_paragraph(); run = par.add_run()
    for tag, attrs in [
        ("w:fldChar",   {"w:fldCharType": "begin"}),
        ("w:instrText", None),
        ("w:fldChar",   {"w:fldCharType": "separate"}),
        ("w:fldChar",   {"w:fldCharType": "end"}),
    ]:
        el = OxmlElement(tag)
        if attrs is None:
            el.set(qn("xml:space"), "preserve"); el.text = 'TOC \\o "1-3" \\h \\z \\u'
        else:
            for k, v in attrs.items(): el.set(qn(k), v)
        run._r.append(el)

# ── helpers de contenido ──────────────────────────────────────────────────────

def h1(text):
    par = doc.add_heading(text, level=1)
    for r in par.runs: r.font.color.rgb = rgb(C_GREEN_DARK); r.font.size = Pt(16)

def h2(text):
    par = doc.add_heading(text, level=2)
    for r in par.runs: r.font.color.rgb = rgb(C_GREEN_MID); r.font.size = Pt(13)

def h3(text):
    par = doc.add_heading(text, level=3)
    for r in par.runs: r.font.color.rgb = rgb(C_GREEN_MID); r.font.size = Pt(11)

def p(text, bold=False, italic=False, center=False, size=11):
    par = doc.add_paragraph()
    if center: par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = par.add_run(text); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    return par

def nota(text):
    """Caja de nota/advertencia en naranja."""
    par = doc.add_paragraph()
    r = par.add_run("Nota: " + text)
    r.italic = True; r.font.size = Pt(10); r.font.color.rgb = rgb(C_AMBER)
    return par

def ph(text):
    doc.add_paragraph()
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = par.add_run("[ " + text + " ]")
    r.italic = True; r.font.size = Pt(10); r.font.color.rgb = rgb(C_GRAY)
    doc.add_paragraph()

def blt(text):
    par = doc.add_paragraph(style="List Bullet")
    r = par.add_run(text); r.font.size = Pt(10)

def tabla(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        shade_cell(cell, C_GREEN_LIGHT); valign_center(cell); cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = rgb(C_GREEN_TEXT)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]; cell.text = str(val)
            for par in cell.paragraphs:
                for r in par.runs: r.font.size = Pt(9)
    doc.add_paragraph()
    return t

# ══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════════

# Universidad y asignatura — parte superior
for _ in range(4):
    doc.add_paragraph()

par_uni = doc.add_paragraph(); par_uni.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_uni = par_uni.add_run("Universidad Europea de Madrid")
r_uni.font.size = Pt(13)

par_esc = doc.add_paragraph(); par_esc.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_esc = par_esc.add_run("Escuela de Arquitectura, Ingeniería, Ciencia y Computación")
r_esc.font.size = Pt(11)

par_grado = doc.add_paragraph(); par_grado.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_grado = par_grado.add_run("Grado en Ingeniería Matemática Aplicada al Análisis de Datos")
r_grado.font.size = Pt(11)

par_asig = doc.add_paragraph(); par_asig.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_asig = par_asig.add_run("Big Data — Documento de seguimiento 2")
r_asig.font.size = Pt(11)

for _ in range(5):
    doc.add_paragraph()

# Título del proyecto
par_tit = doc.add_paragraph(); par_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_tit = par_tit.add_run("Green Turning Point (GTP)")
r_tit.bold = True; r_tit.font.size = Pt(22)

par_sub = doc.add_paragraph(); par_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_sub = par_sub.add_run("Cuadros de mando, KPIs y estructura analítica")
r_sub.font.size = Pt(13); r_sub.italic = True

for _ in range(8):
    doc.add_paragraph()

# Autores y profesor
par_aut = doc.add_paragraph(); par_aut.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_aut = par_aut.add_run(
    "Juan Manuel Palencia Osorio\n"
    "Pablo Mata Rius\n"
    "Pablo Sánchez Ruiz\n"
    "María Paula Aguirre Palacio"
)
r_aut.font.size = Pt(11)

doc.add_paragraph()

par_prof = doc.add_paragraph(); par_prof.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_prof = par_prof.add_run("Profesor: José A. Lozano")
r_prof.font.size = Pt(11)

par_fecha = doc.add_paragraph(); par_fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_fecha = par_fecha.add_run("Abril de 2026")
r_fecha.font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ÍNDICE
# ══════════════════════════════════════════════════════════════════════════════
h1("Índice")
p("Para actualizar el índice en Word: seleccionar todo (Ctrl+A) y pulsar F9.", italic=True, size=10)
doc.add_paragraph(); toc()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. RESUMEN EJECUTIVO
# ══════════════════════════════════════════════════════════════════════════════
h1("1. Resumen Ejecutivo")
p(
    "Green Turning Point nace de una pregunta concreta: ¿se puede saber cuándo una ciudad deja de "
    "contaminar más y empieza a recuperarse? La respuesta está en la Curva de Kuznets Ambiental "
    "(EKC), una hipótesis que sostiene que la relación entre riqueza y degradación ambiental dibuja "
    "una U invertida. A medida que una ciudad se enriquece, primero contamina más; pero llegado un "
    "umbral de renta crítico —el turning point— la tendencia se invierte."
)
p(
    "Hemos convertido esa teoría en un sistema de datos. Tomamos 302 ciudades europeas (Functional "
    "Urban Areas), cruzamos imágenes de satélite con datos macroeconómicos y de política ambiental, "
    "y sobre ese cruce aplicamos cuatro modelos de machine learning en cadena. El resultado es un "
    "ranking de ciudades por su proximidad al turning point, traducido en señales de inversión verde "
    "accesibles a través de una plataforma web y cuadros de mando Power BI."
)
h2("Ecuación central del modelo EKC")
p("ln(Eᵢₜ) = α + β₁·ln(Yᵢₜ) + β₂·[ln(Yᵢₜ)]² + μᵢ + λₜ + εᵢₜ", bold=True)
p("Turning Point:   Y* = exp(−β₁ / 2β₂)", bold=True)
p(
    "Donde Eᵢₜ es el indicador ambiental (NDVI o NO₂) de la ciudad i en el año t, "
    "Yᵢₜ es el PIB per cápita en PPS, μᵢ captura diferencias estructurales entre ciudades "
    "y λₜ absorbe shocks temporales comunes."
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. IDEA DE NEGOCIO Y PLATAFORMA WEB
# ══════════════════════════════════════════════════════════════════════════════
h1("2. Idea de Negocio y Plataforma Web")

h2("2.1 Propuesta de valor")
p(
    "GTP es una plataforma de inversión sostenible que convierte datos satelitales, "
    "macroeconómicos y de política ambiental en señales de inversión concretas sobre 302 ciudades "
    "europeas. Permite a fondos ESG, instituciones públicas y planificadores urbanos identificar "
    "qué ciudades están más próximas a su turning point ecológico y, por tanto, representan "
    "oportunidades de inversión verde con mayor potencial de revalorización."
)

h2("2.2 Estructura de la plataforma web")
p("Diseño minimalista en blanco, gris claro y verde esmeralda. Disponible en español/inglés con soporte para modo oscuro.")
tabla(
    ["Página", "Ruta", "Qué muestra"],
    [
        ["Inicio", "/",
         "Hero con propuesta de valor, métricas de credibilidad, 4 servicios, CTA de registro"],
        ["Servicios", "/servicios",
         "6 tarjetas: Análisis IA, Dashboards Urbanos, Puntuación de Riesgo, Mercados, Alertas, Optimización"],
        ["Dashboard Urbano", "/dashboard-urbano",
         "Filtros Ciudad/Año, 4 KPIs ambientales, gráfico curva EKC, PIB vs NO₂"],
        ["Dashboard Inversión", "/dashboard-inversion",
         "Métricas de portafolio, ranking de ciudades, señal Buy/Hold/Avoid"],
        ["Precios", "/precios",
         "Explorer ($29/mes)  |  Investor ($99/mes)  |  Institutional (custom)"],
        ["Auth", "/auth", "Login y registro por correo y contraseña"],
    ]
)

h2("2.3 Hitos y entregables")
tabla(
    ["#", "Hito", "Descripción", "Estado"],
    [
        ["1", "ETL completo (12 fuentes)",    "Scripts individuales + orquestador run_all.py",     "Completado"],
        ["2", "Pipeline de datos en capas",    "Bronze → Silver → Gold sobre almacenamiento distribuido", "Completado"],
        ["3", "4 modelos ML",                 "K-Means, EKC Regression, XGBoost, Prophet",         "Completado"],
        ["4", "Base de datos para dashboards",  "5 tablas + 5 vistas en MariaDB para Power BI",      "Completado"],
        ["5", "Plataforma web",               "6 secciones (Inicio, Servicios, Dashboards, Auth)", "Completado"],
        ["6", "EDA completo",                 "Análisis exploratorio de las 12 fuentes",           "Completado"],
        ["7", "CMDs operativos internos (Power BI)", "CMD 1 con mockup; CMD 2 planteado teóricamente",    "En proceso"],
        ["8", "CMD cliente ambiental (Power BI)",  "CMD 3 planteado teóricamente",                      "Planteado"],
        ["9", "CMD 4 financiero cliente (Power BI)", "Mockup diseñado; ubicación del mapa pendiente",     "En proceso"],
        ["9", "Prueba end-to-end Lorca",      "Pipeline completo en cluster distribuido UEM",      "Pendiente"],
    ]
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. KPIs Y OKRs
# ══════════════════════════════════════════════════════════════════════════════
h1("3. KPIs y OKRs")
p(
    "Los indicadores clave (KPIs) se organizan en tres dimensiones que reflejan la cadena de valor "
    "del sistema: ambiental, económico-política e inversora. Se han seleccionado 12 KPIs de alto "
    "impacto, 4 por dimensión, evitando duplicidades con variables de ingesta que no aportan valor "
    "de decisión directo al usuario final."
)
p(
    "Los OKRs complementan los KPIs con las aspiraciones estratégicas del proyecto: hacia dónde "
    "apunta el sistema y qué señales indicarían que vamos en la dirección correcta."
)

h2("3.1 KPIs Ambientales")
p(
    "Miden el estado y la evolución de las variables ambientales urbanas que alimentan directamente "
    "la ecuación EKC. Son el punto de partida de toda la cadena analítica y los que el usuario ve "
    "en el Dashboard Urbano de la plataforma web (CMD 3 internamente)."
)
tabla(
    ["KPI", "Fuente", "Columna en BD", "Por qué es relevante"],
    [
        ["NDVI medio",
         "Sentinel-2 (GEE)", "ndvi_mean",
         "Variable E principal de la EKC. Mide el verdor urbano desde satélite. "
         "Un NDVI creciente indica recuperación vegetal."],
        ["NO₂ medio troposférico",
         "Sentinel-5P (GEE)", "no2_mean",
         "Indicador de contaminación complementario. Permite validar la EKC "
         "desde el lado de la calidad del aire."],
        ["Cobertura forestal urbana",
         "ESA WorldCover", "wc_tree_pct",
         "Proxy de reforestación: % de píxeles forestales en el buffer de 10 km de cada ciudad."],
        ["% ciudades por fase",
         "Modelo XGBoost", "turning_point_phase",
         "KPI de resultado: distribución DEGRADANDO / TURNING / RECUPERANDO "
         "sobre las 302 FUAs. El más visible del producto."],
    ]
)

h2("3.2 KPIs Económico-Políticos")
p(
    "Capturan las variables que, según la teoría EKC, determinan cuándo y si una ciudad alcanza "
    "su punto de inflexión. Son la base argumentativa del Dashboard Analítico (CMD 2) y el núcleo "
    "de la propuesta de valor: no solo qué pasa, sino cuánto falta para que cambie."
)
tabla(
    ["KPI", "Fuente", "Columna en BD", "Por qué es relevante"],
    [
        ["PIB per cápita PPS",
         "Eurostat", "gdp_pps_per_capita",
         "Variable Y de la EKC. En Purchasing Power Standards para comparar "
         "ciudades de distintos países con niveles de precios diferentes."],
        ["Y* — Turning Point EKC",
         "Modelo EKC (regresión panel)", "ekc_turning_point_y",
         "Umbral de PIB en el que la degradación ambiental empieza a revertirse. "
         "Output teórico central del modelo."],
        ["GDP gap to turning",
         "Modelo EKC", "gdp_gap_to_turning",
         "Distancia en EUR entre el PIB actual y Y*. Cuanto menor, más próxima "
         "al turning. Señal de inversión directamente accionable."],
        ["EPS Index (0–6)",
         "OECD SDMX", "eps_index",
         "Environmental Policy Stringency: dureza de la regulación ambiental. "
         "Una política más estricta tiende a acelerar el turning point."],
    ]
)

h2("3.3 KPIs de Modelos e Inversión")
p(
    "Son los outputs finales del pipeline de ML que traducen el análisis EKC en señales accionables "
    "para el inversor. Aparecen en el Dashboard de Inversión de la plataforma web y en el CMD 4."
)
tabla(
    ["KPI", "Modelo origen", "Columna en BD", "Por qué es relevante"],
    [
        ["Investment Score (0–100)",
         "Gold Layer (score compuesto)", "investment_score",
         "Puntuación agregada que combina proximidad al turning, tendencia NDVI "
         "y contexto regulatorio. Base del ranking de ciudades."],
        ["Investment Recommendation",
         "Gold Layer", "investment_recommendation",
         "Señal discreta: STRONG_BUY / BUY / HOLD / AVOID. Facilita la "
         "decisión sin necesidad de interpretar el score raw."],
        ["Forecast NDVI 1/3/5 años",
         "Prophet (series temporales)", "prophet_ndvi_forecast_1y/3y/5y",
         "Proyección del NDVI con intervalos de confianza al 95%. "
         "Define el horizonte temporal concreto de la inversión."],
        ["Año estimado de turning point",
         "Prophet", "prophet_turning_year",
         "Año en que el forecast predice la reversión de la tendencia NDVI. "
         "Da al inversor un horizonte temporal concreto."],
    ]
)

h2("3.4 OKRs")
p(
    "Los OKRs recogen las aspiraciones estratégicas del proyecto: hacia dónde apunta el sistema "
    "y qué señales indicarían que vamos en la dirección correcta. No son compromisos cerrados ni "
    "umbrales de aprobación, sino horizontes de referencia que orientan las decisiones de diseño "
    "y permiten evaluar el progreso de forma objetiva. Dado que el resultado de los modelos "
    "estadísticos sobre datos reales es incierto por naturaleza, los Key Results se plantean "
    "como expectativas razonables, no como garantías."
)

h3("O1 — Encontrar evidencia de la hipótesis EKC en ciudades europeas")
p(
    "La aspiración es que los datos reales de las 302 FUAs muestren, al menos en algún grupo "
    "de ciudades, la forma de U invertida que predice la teoría EKC. Incluso un resultado "
    "parcial aportaría valor académico y orientaría la selección de ciudades para el producto."
)
tabla(
    ["Key Result", "Expectativa orientativa", "Cómo se evaluaría"],
    [
        ["KR1",
         "Al menos 1 cluster con β₂ negativo y Y* en rango económicamente razonable (15k–80k EUR PPS)",
         "ekc_parameters.beta2 < 0  AND  turning_point_y_star BETWEEN 15000 AND 80000"],
        ["KR2",
         "Una proporción significativa de ciudades en fases positivas (TURNING o RECUPERANDO)",
         "Se espera >10% del total en fact_kuznets.turning_point_phase"],
        ["KR3",
         "Algún cluster con capacidad explicativa mínima detectable en el panel",
         "ekc_parameters.r_squared > 0.10 en al menos un cluster"],
    ]
)

h3("O2 — Que el sistema genere señales de inversión útiles y accesibles")
p(
    "Más allá de la validación teórica, la aspiración es que el pipeline produzca outputs "
    "concretos que un inversor pueda consultar sin necesidad de interpretar modelos directamente."
)
tabla(
    ["Key Result", "Expectativa orientativa", "Cómo se evaluaría"],
    [
        ["KR1",
         "Cobertura amplia de ciudades con Investment Score calculado y distribuido de forma heterogénea",
         "Se espera investment_score NOT NULL para ≥200 ciudades"],
        ["KR2",
         "Prophet capaz de generar forecasts para la mayoría de ciudades con series suficientes",
         "Se espera prophet_ndvi_forecast_1y disponible para ≥100 ciudades"],
        ["KR3",
         "Dashboards con tiempos de carga adecuados para uso en tiempo real",
         "Tiempo de carga objetivo <5 s en Power BI sobre MariaDB"],
    ]
)

h3("O3 — Pipeline reproducible y autónomo en el cluster Lorca")
p(
    "La aspiración técnica es que cualquier miembro del equipo pueda lanzar el pipeline "
    "completo en Lorca con un único comando y obtener los mismos resultados. "
    "La reproducibilidad es un requisito de calidad, no un extra."
)
tabla(
    ["Key Result", "Expectativa orientativa", "Cómo se evaluaría"],
    [
        ["KR1",
         "Pipeline end-to-end ejecutado sin intervención manual ni correcciones sobre la marcha",
         "run_pipeline.py finaliza con exit code 0 en el cluster Lorca"],
        ["KR2",
         "Resultados estables entre ejecuciones sucesivas (overwrite idempotente de particiones)",
         "Segunda ejecución produce los mismos outputs que la primera"],
    ]
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. ESTRUCTURA ANALÍTICA
# ══════════════════════════════════════════════════════════════════════════════
h1("4. Estructura Analítica")
p(
    "La estructura analítica describe cómo están organizados los datos desde la ingesta hasta el "
    "servicio final al usuario. El proyecto sigue la arquitectura Medallion —tres capas "
    "progresivas (Bronze, Silver, Gold)— sobre HDFS, el sistema de almacenamiento distribuido "
    "del cluster. Spark actúa como motor de procesamiento entre capas, e Hive proporciona "
    "el catálogo de tablas que estructura los datos. La salida final va a MariaDB, "
    "que alimenta directamente los cuadros de mando Power BI."
)

h2("4.1 Flujo completo de datos")
p("El dato recorre seis etapas desde las fuentes externas hasta el dashboard del usuario:")
p(
    "Fuentes externas (12 APIs / datasets)  →  ETL Python  →  CSVs  →  "
    "Bronze (datos en bruto, 12 tablas en HDFS)  →  Silver (esquema en estrella de Kimball, Hive)  →  "
    "Gold (resultados analíticos + ML, Hive)  →  MariaDB  →  Power BI / Web"
)
try:
    doc.add_picture(IMG, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
except Exception:
    ph("IMAGEN: Diagrama de arquitectura Medallion")

h2("4.2 Capa Bronze — Datos en bruto (12 tablas)")
p(
    "Primera capa de almacenamiento: los datos llegan de cada fuente y se guardan tal cual, "
    "sin transformaciones. Se almacenan en formato Parquet con compresión Snappy —columnar "
    "y comprimido— lo que permite consultas rápidas sobre grandes volúmenes. Cada tabla "
    "está particionada por año y país, lo que acelera los filtros temporales y geográficos "
    "de los cuadros de mando. Hive gestiona el catálogo de estas tablas sobre HDFS."
)
tabla(
    ["Tabla Bronze (Hive)", "Fuente", "Variables clave"],
    [
        ["bronze_sentinel2",    "GEE Sentinel-2",     "ndvi_mean, ndvi_std, ndvi_p10, ndvi_p90"],
        ["bronze_s5p",          "GEE Sentinel-5P",    "no2_mean, no2_max, no2_std"],
        ["bronze_hrl",          "Copernicus HRL",     "imperviousness_mean, tree_cover_pct"],
        ["bronze_finance",      "Yahoo Finance",      "ticker, close_price, annual_volatility"],
        ["bronze_eurostat",     "Eurostat",           "gdp_pps_per_capita, fua_population"],
        ["bronze_era5",         "GEE ERA5-Land",      "temp_annual_mean_c, precip_annual_sum_m"],
        ["bronze_s5p_aerosol",  "GEE S5P UVAI",       "uvai_annual_mean"],
        ["bronze_worldcover",   "GEE ESA WorldCover", "wc_tree_pct, wc_built_pct, wc_crop_pct"],
        ["bronze_edgar_co2",    "EDGAR v8",           "co2_country_kt"],
        ["bronze_oecd_raw",     "OECD SDMX",          "eps_index, env_tax_usd, ghg_total_kt"],
        ["bronze_investeu_raw", "EIB / InvestEU",     "investeu_ops_count, investeu_total_eur"],
        ["bronze_merge",        "Derivado (merge.py)","ndvi_slope, pct_green_area"],
    ]
)

h2("4.3 Capa Silver — Esquema en estrella de Kimball")
p(
    "Segunda capa: los datos en bruto de las 12 fuentes se limpian, homogenizan y se integran "
    "en un esquema en estrella de Kimball. Spark SQL realiza estas transformaciones sobre HDFS "
    "y Hive registra la estructura resultante. El esquema organiza la información en tablas de "
    "contexto (dimensiones) y tablas de medidas (hechos), facilitando las consultas analíticas "
    "de los cuadros de mando."
)
tabla(
    ["Tabla Silver", "Tipo", "Descripción"],
    [
        ["dim_city",                  "Contexto (dimensión)",
         "Las 302 ciudades europeas: nombre, país y coordenadas geográficas. "
         "Guarda historial de cambios (SCD-2) para rastrear actualizaciones."],
        ["dim_date",                  "Contexto (dimensión)",
         "Calendario completo: año, mes, trimestre."],
        ["dim_cluster",               "Contexto (dimensión)",
         "Los 4 grupos de ciudades del modelo K-Means con su etiqueta y parámetros EKC."],
        ["dim_company",               "Contexto (dimensión)",
         "43 empresas del sector verde europeo con su sector e industria. "
         "También con historial de cambios (SCD-2)."],
        ["fact_kuznets",              "Medidas (hecho central)",
         "Indicadores ambientales, económicos y resultados ML por ciudad y año."],
        ["fact_finance",              "Medidas (hecho)",
         "Serie histórica bursátil: precio, volatilidad y retorno anual."],
        ["fact_environmental_policy", "Medidas (hecho)",
         "Política ambiental e inversión verde por país y año (OECD, InvestEU, EDGAR)."],
    ]
)

h2("4.4 Capa Gold — Resultados analíticos")
p(
    "Tercera capa: contiene los resultados finales de los cuatro modelos de machine learning "
    "aplicados sobre Silver mediante Spark. Las tablas se almacenan en HDFS bajo catálogo "
    "Hive y desde ahí se exportan a MariaDB para los cuadros de mando."
)
tabla(
    ["Tabla Gold (Hive)", "Granularidad", "Para qué sirve"],
    [
        ["fact_kuznets",   "Ciudad × Año",       "Tabla central: todos los indicadores + outputs de los 4 modelos."],
        ["city_ranking",   "Ciudad × Año",       "Ranking anual por investment_score. Alimenta CMD 1 y CMD 4."],
        ["ekc_parameters", "Cluster × Estimación","Parámetros β₁, β₂, Y*, R², forma de curva. Alimenta CMD 2."],
        ["model_results",  "Ciudad × Año",       "Outputs detallados: fases XGBoost, forecasts Prophet, confianza."],
    ]
)

h2("4.5 Base de datos para los cuadros de mando (MariaDB)")
p(
    "Las tablas Gold se exportan a MariaDB, que actúa como la fuente de datos que Power BI "
    "consulta directamente. Se crean además 5 vistas con los datos ya preparados para cada "
    "tipo de cuadro de mando, de forma que Power BI no necesita hacer cálculos adicionales."
)
p(
    "El prototipo interactivo de la plataforma web está disponible en: "
    "https://green-turning-point.lovable.app"
)
tabla(
    ["Tabla / Vista MariaDB", "Para qué sirve en los CMDs"],
    [
        ["dim_city",                "Coordenadas para el mapa europeo (CMD 1)"],
        ["fact_kuznets",            "Métricas ciudad × año — base de CMD 2 y CMD 3"],
        ["city_ranking",            "Ranking de ciudades por año — CMD 1 y CMD 4"],
        ["ekc_parameters",          "Parámetros EKC por cluster — CMD 2"],
        ["model_results",           "Outputs ML detallados — CMD 3 y CMD 4"],
        ["v_powerbi_map",           "Vista para mapa coroplético: lat, lon, investment_score, fase — CMD 1"],
        ["v_turning_opportunities", "Ciudades en fase TURNING activa — CMD 4"],
        ["v_ekc_summary",           "Resumen EKC por cluster — CMD 2"],
        ["v_top_20_cities",         "Top 20 ciudades por año — CMD 1"],
        ["v_phase_distribution",    "Distribución de fases por país y año — CMD 3"],
    ]
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. CUADROS DE MANDO
# ══════════════════════════════════════════════════════════════════════════════
h1("5. Cuadros de Mando")
p(
    "Se diseñan cuatro cuadros de mando en Power BI. Los dos primeros (CMD 1 y CMD 2) son "
    "operativos e internos, orientados al equipo analítico y de validación. El tercero (CMD 3) "
    "es el cuadro de mando de cliente, orientado al usuario final de la plataforma. El cuarto "
    "(CMD 4) está en fase de clasificación: su contenido de inversión lo sitúa como posible "
    "cuadro de cliente, pero también podría mantenerse como complemento interno del CMD 1. "
    "Todos se alimentan de las vistas de MariaDB conectadas directamente a Power BI."
)

# ──────────────────────────────────────────────────────────────────────────────
h2("5.1  Cuadros de Mando Operativos  (internos)")
# ──────────────────────────────────────────────────────────────────────────────
p(
    "Orientados al equipo analítico, investigadores y al evaluador. Su función es visualizar "
    "el mapa de oportunidades de inversión a nivel interno y validar los resultados del modelo "
    "EKC con los datos reales del pipeline. CMD 1 cuenta con mockup interactivo desarrollado "
    "en Lovable; CMD 2 está planteado teóricamente."
)

# ── CMD 1 ──
h3("CMD 1 — Mapa Europeo de Inversión  ·  Vista Interna")
p(
    "Audiencia interna: dirección, equipo analítico, comités de validación. "
    "Vista principal: v_powerbi_map. "
    "Permite ver de un vistazo qué ciudades europeas ofrecen las mejores oportunidades de "
    "inversión verde según el modelo, y sirve como herramienta de validación del pipeline."
)
p(
    "El elemento central es un mapa coroplético de Europa donde cada una de las 302 ciudades "
    "aparece coloreada del rojo al verde según su investment_score. Al hacer clic en una ciudad "
    "se despliegan sus KPIs en las tarjetas laterales."
)
p(
    "Prototipo interactivo funcional disponible en: https://green-turning-point.lovable.app  "
    "(sección Dashboard Inversión → vista ejecutiva)"
)
nota(
    "Las capturas siguientes proceden del prototipo Lovable. "
    "El mapa coroplético de Europa está planteado en el diseño pero aún no implementado "
    "en la versión actual del prototipo; se incluirá en la conexión final con MariaDB."
)
tabla(
    ["Componente", "Qué muestra", "Vista / Tabla MariaDB"],
    [
        ["Mapa coroplético europeo",
         "302 ciudades coloreadas por investment_score (rojo → verde). Clic activa filtros.",
         "v_powerbi_map (latitude, longitude, investment_score, turning_point_phase)"],
        ["Filtros interactivos",
         "Año, país, cluster EKC, fase (DEGRADANDO / TURNING / RECUPERANDO)",
         "fact_kuznets"],
        ["Top 20 ciudades",
         "Ranking con nombre, país, fase e investment_score",
         "v_top_20_cities"],
        ["Semáforo de recomendación",
         "STRONG_BUY / BUY / HOLD / AVOID para la ciudad seleccionada",
         "fact_kuznets.investment_recommendation"],
        ["4 KPI cards",
         "N.º ciudades TURNING  |  N.º RECUPERANDO  |  Y* medio  |  Score medio",
         "fact_kuznets + ekc_parameters"],
    ]
)
ph("INSERTAR CAPTURA CMD 1 — vista principal con layout del mapa y KPI cards")
ph("INSERTAR CAPTURA CMD 1 — detalle de filtros y tabla Top 20 ciudades")

# ── CMD 2 ──
h3("CMD 2 — Dashboard EKC  ·  Perfil Analítico")
p(
    "Audiencia: economistas, académicos, equipo técnico. "
    "Permite explorar la curva de Kuznets por cluster y validar la hipótesis EKC "
    "con los datos reales del pipeline."
)
nota(
    "CMD 2 está planteado teóricamente: se definen sus componentes, fuentes de datos y propósito, "
    "pero no se ha desarrollado mockup visual. La implementación en Power BI se realizará "
    "en fases posteriores del proyecto."
)
tabla(
    ["Componente", "Qué muestra", "Vista / Tabla MariaDB"],
    [
        ["Filtros",
         "Cluster, ciudad, país, rango de años",
         "fact_kuznets"],
        ["Curva EKC por cluster (scatter)",
         "GDP per cápita vs NDVI con parábola ajustada (β₁, β₂ del cluster)",
         "ekc_parameters + fact_kuznets"],
        ["Tabla parámetros EKC",
         "β₁, β₂, Y*, R² ajustado, forma de curva (INVERTED_U / MONOTONIC)",
         "v_ekc_summary"],
        ["GDP gap to turning (barras)",
         "Cuánto PIB per cápita falta para el turning point por ciudad",
         "fact_kuznets.gdp_gap_to_turning"],
        ["Mapa turning year Prophet",
         "Año estimado de inflexión por ciudad sobre mapa europeo",
         "model_results.prophet_turning_year + dim_city"],
    ]
)

# ──────────────────────────────────────────────────────────────────────────────
h2("5.2  Cuadros de Mando de Cliente  (externos)")
# ──────────────────────────────────────────────────────────────────────────────
p(
    "Orientados al usuario final de la plataforma: inversores, fondos ESG, analistas ambientales "
    "y planificadores urbanos. CMD 3 está planteado teóricamente. CMD 4 cuenta con mockup "
    "interactivo desarrollado en Lovable."
)
nota(
    "El mapa coroplético europeo de inversión está pendiente de ubicación definitiva: "
    "puede incorporarse como visualización independiente dentro de CMD 4 para evitar scroll, "
    "o bien como componente complementario integrado en el propio CMD 4. "
    "Ambas opciones se valorarán en la fase de implementación en Power BI."
)

# ── CMD 3 ──
h3("CMD 3 — Dashboard Ambiental  ·  Perfil Operativo")
p(
    "Audiencia: analistas ambientales, investigadores, planificadores urbanos. "
    "Conecta conceptualmente con la sección /dashboard-urbano de la plataforma web. "
    "Permite monitorizar la evolución ambiental de cada ciudad en detalle."
)
nota(
    "CMD 3 está planteado teóricamente: se definen sus componentes, fuentes de datos y propósito, "
    "pero no se ha desarrollado mockup visual. La implementación en Power BI se realizará "
    "en fases posteriores del proyecto."
)
tabla(
    ["Componente", "Qué muestra", "Vista / Tabla MariaDB"],
    [
        ["Serie temporal NDVI",
         "Evolución anual del NDVI por ciudad / país de 2018 a 2024",
         "fact_kuznets (ndvi_mean, year, city_code)"],
        ["Mapa calor NDVI slope",
         "Ciudades del país seleccionado que mejoran vs. que empeoran en vegetación",
         "model_results (ndvi_trend_slope) + dim_city"],
        ["PIB vs NO₂ (scatter)",
         "Relación entre riqueza y contaminación por ciudad y año",
         "fact_kuznets (gdp_pps_per_capita, no2_mean)"],
        ["4 KPI cards",
         "NDVI actual  |  NO₂ actual  |  Cobertura forestal  |  PIB per cápita",
         "fact_kuznets (último año disponible)"],
        ["Filtros",
         "Ciudad, país, año (2018–2024)",
         "fact_kuznets"],
    ]
)

# ──────────────────────────────────────────────────────────────────────────────
h3("CMD 4 — Dashboard Financiero  ·  Perfil Inversor")
p(
    "Audiencia: gestores de cartera, fondos ESG, inversores sostenibles. "
    "Conecta con la sección /dashboard-inversión de la plataforma web. "
    "Traduce los resultados del modelo en señales de inversión accionables."
)
p(
    "Prototipo interactivo funcional disponible en: https://green-turning-point.lovable.app  "
    "(sección Dashboard Inversión)"
)
tabla(
    ["Componente", "Qué muestra", "Vista / Tabla MariaDB"],
    [
        ["Ranking de oportunidades",
         "Ciudades con señal Buy/Hold/Avoid ordenadas por investment_score",
         "v_turning_opportunities + city_ranking"],
        ["Scatter inversión vs riesgo",
         "Investment score (eje Y) vs volatilidad del sector verde (eje X) por ciudad",
         "city_ranking + fact_finance (fin_annual_volatility)"],
        ["Volatilidad sectorial",
         "Ciclos de riesgo del sector verde europeo por año",
         "fact_finance (fin_annual_volatility, year)"],
        ["4 KPI cards portafolio",
         "Score medio  |  Retorno estimado  |  Diversificación geográfica  |  Riesgo",
         "Calculados en Power BI sobre city_ranking + model_results"],
    ]
)
ph("INSERTAR CAPTURA CMD 4 — vista principal con ranking y scatter inversión vs riesgo")
ph("INSERTAR CAPTURA CMD 4 — detalle de KPI cards de portafolio")

# ══════════════════════════════════════════════════════════════════════════════
# PIE DE PÁGINA + GUARDAR
# ══════════════════════════════════════════════════════════════════════════════
add_page_numbers()
doc.save(DST)
print("Documento guardado: " + DST)
if os.path.exists(IMG):
    os.remove(IMG)
